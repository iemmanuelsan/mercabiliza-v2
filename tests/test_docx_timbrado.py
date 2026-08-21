"""Testes do formulário DOCX e do papel timbrado."""

from __future__ import annotations

import io

import pytest
from docx import Document

from src.core.models import Endereco
from src.core.pessoas import ContratantePF, ContratantePJ, RepresentanteLegal
from src.exporters.docx_abertura import (
    MARCADOR_PENDENTE,
    dados_de_contratante,
    gerar_formulario_abertura,
)
from src.exporters.timbrado import TIMBRADO_PADRAO, Timbrado, timbrado_padrao

END = Endereco(logradouro="RUA TENENTE RUI", numero="120", bairro="Centro",
               municipio="Macaé", uf="RJ", cep="27910000")

MEI = ContratantePJ(
    razao_social="ANDERSON ANDRADE MONTEIRO", cnpj="63435477000110",
    nome_fantasia="EasyMarket Molinere", regime="MEI",
    natureza_juridica="Empresário (Individual)", data_abertura="30/10/2025",
    cnae_principal="4723700 - Comércio varejista de bebidas", endereco=END,
    telefone="(22) 99999-1234", email="contato@easymarket.com.br",
    representante=RepresentanteLegal(
        nome="Anderson Andrade Monteiro", cpf="52998224725", rg="12.345.678",
        orgao_emissor="DETRAN/RJ", estado_civil="casado", profissao="empresário"),
)

PF = ContratantePF(nome="Vinicius Almeida", cpf="11144477735", endereco=END,
                   telefone="(22) 98888-7777", email="v@easymarket.com.br")


def _texto(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            partes.extend(c.text for c in row.cells)
    return "\n".join(partes)


def _negritos(docx_bytes: bytes) -> set[str]:
    """Textos marcados em negrito — a convenção de 'preenchido pelo sistema'."""
    doc = Document(io.BytesIO(docx_bytes))
    achados: set[str] = set()
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        if r.bold and r.text.strip():
                            achados.add(r.text.strip())
    return achados


# --------------------------------------------------------------------------- #
# DOCX                                                                         #
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize("perfil", ["PF", "MEI", "PJ"])
def test_docx_gera_arquivo_valido(perfil):
    empresa, endereco, socios = dados_de_contratante(MEI)
    dados = gerar_formulario_abertura(perfil, empresa, endereco, socios)
    assert dados[:2] == b"PK"          # container zip do OOXML
    assert len(dados) > 10_000
    Document(io.BytesIO(dados))        # abre sem erro


def test_titulo_muda_por_perfil():
    empresa, endereco, socios = dados_de_contratante(MEI)
    assert "DESENQUADRAMENTO" in _texto(
        gerar_formulario_abertura("MEI", empresa, endereco, socios))
    assert "ABERTURA DE EMPRESA" in _texto(
        gerar_formulario_abertura("PF", empresa, endereco, socios))
    assert "ALTERAÇÃO CONTRATUAL" in _texto(
        gerar_formulario_abertura("PJ", empresa, endereco, socios))


def test_dados_conhecidos_saem_em_negrito():
    """O que veio da API não deve pedir digitação de novo."""
    empresa, endereco, socios = dados_de_contratante(MEI)
    docx = gerar_formulario_abertura("MEI", empresa, endereco, socios)
    negritos = _negritos(docx)
    assert "RUA TENENTE RUI" in negritos
    assert "Macaé" in negritos
    assert "529.982.247-25" in negritos, "CPF deve sair com máscara"


def test_decisoes_do_cliente_saem_marcadas():
    empresa, endereco, socios = dados_de_contratante(MEI)
    texto = _texto(gerar_formulario_abertura("MEI", empresa, endereco, socios))
    assert MARCADOR_PENDENTE in texto
    # Estes são os campos que só o cliente decide.
    for rotulo in ("2ª opção de Razão Social", "Capital Social",
                   "Participação no capital", "Sócio administrador"):
        assert rotulo in texto


def test_perfil_mei_tem_bloco_de_desenquadramento():
    empresa, endereco, socios = dados_de_contratante(MEI)
    desenq = {"cnpj": "63.435.477/0001-10", "razao_atual": "ANDERSON",
              "abertura": "30/10/2025"}
    com = _texto(gerar_formulario_abertura("MEI", empresa, endereco, socios, desenq))
    sem = _texto(gerar_formulario_abertura("PF", empresa, endereco, socios))
    assert "DESENQUADRAMENTO DO MEI" in com
    assert "Faturamento acumulado no ano" in com
    assert "DESENQUADRAMENTO DO MEI" not in sem


def test_sempre_imprime_blocos_de_socio_para_o_cliente_completar():
    """Mesmo conhecendo 1 sócio, sai espaço para o segundo — o cliente pode
    incluir alguém que a contabilidade ainda não conhece."""
    _, endereco, socios = dados_de_contratante(MEI)
    texto = _texto(gerar_formulario_abertura("PF", {}, endereco, socios,
                                             minimo_socios=2))
    assert "SÓCIO 01" in texto
    assert "SÓCIO 02" in texto


def test_pf_nao_traz_dados_de_empresa_inexistente():
    _, endereco, socios = dados_de_contratante(PF)
    texto = _texto(gerar_formulario_abertura("PF", {}, endereco, socios))
    assert "Razão Social (1ª opção)" in texto
    assert "ANDERSON" not in texto


def test_documentos_a_anexar_variam_por_perfil():
    empresa, endereco, socios = dados_de_contratante(MEI)
    mei = _texto(gerar_formulario_abertura("MEI", empresa, endereco, socios))
    pf = _texto(gerar_formulario_abertura("PF", empresa, endereco, socios))
    assert "CCMEI" in mei
    assert "CCMEI" not in pf


def test_contratante_vazio_nao_quebra():
    """Usuário clica em gerar sem preencher nada."""
    empresa, endereco, socios = dados_de_contratante(ContratantePF())
    assert gerar_formulario_abertura("PF", empresa, endereco, socios)[:2] == b"PK"


# --------------------------------------------------------------------------- #
# Timbrado                                                                     #
# --------------------------------------------------------------------------- #
def test_timbrado_oficial_esta_versionado():
    """Se o PNG não subir para o repositório, o contrato sai sem a marca."""
    assert TIMBRADO_PADRAO.exists(), f"timbrado ausente em {TIMBRADO_PADRAO}"


def test_timbrado_padrao_usa_a_imagem_quando_existe():
    assert timbrado_padrao(usar_timbrado=True).tem_imagem


def test_timbrado_desligado_cai_no_fallback():
    """Minuta interna sai sem a arte, gastando menos tinta."""
    assert not timbrado_padrao(usar_timbrado=False).tem_imagem


def test_margem_superior_reserva_espaco_para_o_logo():
    """Se a margem encolher, o texto sobe e invade o logo impresso."""
    assert timbrado_padrao(True).margem_topo_mm >= 30


def test_imagem_inexistente_nao_quebra_o_pdf():
    from pathlib import Path

    from src.core.contrato import ParametrosContrato
    from src.core.pessoas import contratada_padrao
    from src.exporters.pdf_contrato import gerar_contrato

    quebrado = Timbrado(imagem=Path("/nao/existe/timbrado.png"))
    assert not quebrado.tem_imagem
    pdf = gerar_contrato(MEI, contratada_padrao(),
                         ParametrosContrato(valor_mensal=350.0, foro="Macaé/RJ"))
    assert pdf.startswith(b"%PDF")

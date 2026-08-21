"""Formulário de Abertura / Desenquadramento em DOCX.

Replica a estrutura da "FICHA CADASTRAL - ABERTURA DE EMPRESA" da Mercabiliza,
com o diferencial que motiva o módulo: **preenchimento híbrido**.

## Por que DOCX e não XLSX

O formulário original é um documento com tabelas rotuladas, não uma planilha
de dados. Em DOCX o cliente preenche no Word, no Google Docs ou imprime e
escreve à mão, sem quebrar layout — e é o formato que ele já conhece. XLSX
faria sentido se o retorno fosse importado de volta em massa; para um
formulário assinado por cliente, DOCX é o certo.

## Convenção visual

* **Negrito** = preenchido pelo sistema (via API de CNPJ, CEP ou cadastro).
  O cliente só confere.
* `[ PREENCHER AQUI ]` em vermelho = depende de decisão do cliente (opções de
  razão social, capital social, distribuição de quotas, previsão de
  faturamento).
* Campo vazio = dado que o cliente tem mas o sistema não conseguiu obter.

Essa distinção é o ponto do documento: sem ela o cliente relê tudo, e com ela
ele vai direto às ~8 decisões que só ele pode tomar.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..config import (
    CONTRATADA_EMAIL,
    CONTRATADA_NOME_FANTASIA,
    CONTRATADA_TELEFONE,
)
from ..core.cnpj import formatar as formatar_cnpj
from ..core.contrato import valor_extenso
from ..core.cpf import formatar as formatar_cpf
from ..core.formatters import moeda

MARCADOR_PENDENTE = "[ PREENCHER AQUI ]"

COR_MARCA = RGBColor(0xDC, 0x32, 0x50)
COR_PENDENTE = RGBColor(0xC0, 0x1B, 0x36)
COR_CINZA = RGBColor(0x55, 0x5B, 0x66)
CINZA_FUNDO = "F5F6F8"


# --------------------------------------------------------------------------- #
# Modelo de campo                                                              #
# --------------------------------------------------------------------------- #
@dataclass(frozen=True, slots=True)
class Campo:
    """Um campo do formulário.

    ``valor`` vazio + ``pendente=True`` → sai como ``[ PREENCHER AQUI ]``.
    ``valor`` preenchido → sai em negrito (veio do sistema).
    ``valor`` vazio + ``pendente=False`` → sai como linha em branco.
    """

    rotulo: str
    valor: str = ""
    pendente: bool = False
    dica: str = ""

    @property
    def preenchido_pelo_sistema(self) -> bool:
        return bool(self.valor.strip())


@dataclass(slots=True)
class Secao:
    titulo: str
    campos: list[Campo] = field(default_factory=list)
    colunas: int = 2


# --------------------------------------------------------------------------- #
# Helpers de formatação                                                        #
# --------------------------------------------------------------------------- #
def _sombrear(celula, cor_hex: str) -> None:
    """Fundo de célula — o python-docx não expõe isso na API pública."""
    elemento = OxmlElement("w:shd")
    elemento.set(qn("w:val"), "clear")
    elemento.set(qn("w:fill"), cor_hex)
    celula._tc.get_or_add_tcPr().append(elemento)


def _run(paragrafo, texto: str, *, negrito: bool = False, tamanho: float = 9,
         cor: RGBColor | None = None, italico: bool = False):
    r = paragrafo.add_run(texto)
    r.bold = negrito
    r.italic = italico
    r.font.size = Pt(tamanho)
    if cor is not None:
        r.font.color.rgb = cor
    return r


def _escrever_campo(celula, campo: Campo) -> None:
    """Rótulo + valor numa célula, aplicando a convenção visual."""
    celula.text = ""
    p = celula.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    _run(p, f"{campo.rotulo}: ", negrito=False, tamanho=8.5, cor=COR_CINZA)

    if campo.preenchido_pelo_sistema:
        _run(p, campo.valor, negrito=True, tamanho=9)
    elif campo.pendente:
        _run(p, MARCADOR_PENDENTE, negrito=True, tamanho=9, cor=COR_PENDENTE)
    else:
        _run(p, "_" * 28, tamanho=9, cor=COR_CINZA)

    if campo.dica:
        p2 = celula.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        _run(p2, campo.dica, tamanho=7, cor=COR_CINZA, italico=True)


def _titulo_secao(doc, texto: str) -> None:
    tabela = doc.add_table(rows=1, cols=1)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    celula = tabela.rows[0].cells[0]
    _sombrear(celula, "2F5597")
    celula.text = ""
    p = celula.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, texto.upper(), negrito=True, tamanho=9.5, cor=RGBColor(0xFF, 0xFF, 0xFF))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _tabela_campos(doc, campos: list[Campo], colunas: int = 2) -> None:
    if not campos:
        return
    linhas = -(-len(campos) // colunas)
    tabela = doc.add_table(rows=linhas, cols=colunas)
    tabela.style = "Table Grid"
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, campo in enumerate(campos):
        celula = tabela.cell(i // colunas, i % colunas)
        _escrever_campo(celula, campo)
        if campo.pendente and not campo.preenchido_pelo_sistema:
            _sombrear(celula, "FDF0F2")

    # Células sobrando na última linha ficam em branco, não com "None".
    for j in range(len(campos), linhas * colunas):
        tabela.cell(j // colunas, j % colunas).text = ""

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# --------------------------------------------------------------------------- #
# Montagem das seções                                                          #
# --------------------------------------------------------------------------- #
def _secao_empresa(dados: dict) -> Secao:
    """Dados da futura empresa (ou da nova configuração, no desenquadramento)."""
    return Secao("Dados da empresa", [
        Campo("Razão Social (1ª opção)", dados.get("razao_social", ""),
              pendente=True, dica="Nome empresarial pretendido"),
        Campo("2ª opção de Razão Social", "", pendente=True),
        Campo("3ª opção de Razão Social", "", pendente=True),
        Campo("Nome Fantasia", dados.get("nome_fantasia", ""), pendente=True),
        Campo("Descrição da Atividade", dados.get("atividade", ""), pendente=True,
              dica="O que a empresa vende/faz, em uma frase"),
        Campo("CNAE pretendido", dados.get("cnae", ""),
              dica="Sugerido pela contabilidade"),
        Campo("Capital Social (R$)", dados.get("capital_social", ""), pendente=True),
        Campo("Capital Social por extenso", dados.get("capital_extenso", ""),
              pendente=True),
        Campo("Tipo Jurídico", dados.get("tipo_juridico", ""), pendente=True,
              dica="Sociedade Limitada, Sociedade Unipessoal, Empresário Individual"),
        Campo("Previsão de faturamento mensal", dados.get("faturamento", ""),
              pendente=True),
        Campo("Regime tributário pretendido", dados.get("regime", ""),
              dica="Sugerido pela contabilidade após análise"),
        Campo("Nº de funcionários previstos", dados.get("funcionarios", ""),
              pendente=True),
    ])


def _secao_endereco(dados: dict) -> Secao:
    return Secao("Endereço da empresa", [
        Campo("Logradouro", dados.get("logradouro", "")),
        Campo("Número", dados.get("numero", "")),
        Campo("Complemento", dados.get("complemento", "")),
        Campo("Bairro", dados.get("bairro", "")),
        Campo("Município", dados.get("municipio", "")),
        Campo("UF", dados.get("uf", "")),
        Campo("CEP", dados.get("cep", "")),
        Campo("Ponto de referência", "", pendente=False),
        Campo("Telefone", dados.get("telefone", "")),
        Campo("E-mail", dados.get("email", "")),
        Campo("O imóvel é próprio ou alugado?", "", pendente=True),
        Campo("IPTU / inscrição imobiliária", "", pendente=True,
              dica="Necessário para o alvará"),
    ])


def _secao_socio(numero: int, dados: dict) -> Secao:
    return Secao(f"Sócio {numero:02d}", [
        Campo("Nome completo", dados.get("nome", "")),
        Campo("Nacionalidade", dados.get("nacionalidade", "brasileiro(a)")),
        Campo("Naturalidade", dados.get("naturalidade", ""), pendente=True),
        Campo("Profissão", dados.get("profissao", "")),
        Campo("Data de nascimento", dados.get("nascimento", ""), pendente=True),
        Campo("Estado civil", dados.get("estado_civil", "")),
        Campo("Regime de bens", dados.get("regime_bens", ""), pendente=True,
              dica="Parcial, Total, Universal — se casado(a)"),
        Campo("CPF/MF", dados.get("cpf", "")),
        Campo("C.I. / R.G.", dados.get("rg", "")),
        Campo("Órgão emissor / UF", dados.get("orgao", "")),
        Campo("Data de expedição", dados.get("expedicao", ""), pendente=True),
        Campo("Título de eleitor", "", pendente=True),
        Campo("Participação no capital (%)", dados.get("participacao", ""),
              pendente=True),
        Campo("Sócio administrador?", dados.get("administrador", ""), pendente=True,
              dica="Sim ou Não"),
        Campo("Logradouro", dados.get("logradouro", "")),
        Campo("Número", dados.get("numero", "")),
        Campo("Complemento", dados.get("complemento", "")),
        Campo("Bairro", dados.get("bairro", "")),
        Campo("Município", dados.get("municipio", "")),
        Campo("UF", dados.get("uf", "")),
        Campo("CEP", dados.get("cep", "")),
        Campo("Telefone / Celular", dados.get("telefone", "")),
        Campo("E-mail", dados.get("email", "")),
    ])


def _secao_desenquadramento(dados: dict) -> Secao:
    """Só aparece no perfil MEI — o que muda ao sair do MEI."""
    return Secao("Desenquadramento do MEI", [
        Campo("CNPJ atual (MEI)", dados.get("cnpj", "")),
        Campo("Razão social atual", dados.get("razao_atual", "")),
        Campo("Data de abertura", dados.get("abertura", "")),
        Campo("Faturamento acumulado no ano", dados.get("faturamento_ano", ""),
              pendente=True, dica="Determina se o desenquadramento é retroativo"),
        Campo("Faturamento do ano anterior", "", pendente=True),
        Campo("Data pretendida de efeito", "", pendente=True,
              dica="Retroativo a 01/01 ou a partir do mês seguinte"),
        Campo("Possui funcionário registrado?", "", pendente=True),
        Campo("Emite NFC-e atualmente?", dados.get("nfce", ""), pendente=True),
        Campo("Sistema de gestão / PDV", dados.get("sistema", ""), pendente=True),
        Campo("Possui inscrição estadual?", dados.get("ie", ""), pendente=True),
    ])


# --------------------------------------------------------------------------- #
# Documento                                                                    #
# --------------------------------------------------------------------------- #
def _configurar_pagina(doc) -> None:
    for secao in doc.sections:
        secao.top_margin = Cm(1.8)
        secao.bottom_margin = Cm(1.8)
        secao.left_margin = Cm(1.8)
        secao.right_margin = Cm(1.8)


def _cabecalho(doc, titulo: str, subtitulo: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, CONTRATADA_NOME_FANTASIA.upper(), negrito=True, tamanho=16, cor=COR_MARCA)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, titulo, negrito=True, tamanho=13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, subtitulo, tamanho=8.5, cor=COR_CINZA)


def _legenda(doc) -> None:
    tabela = doc.add_table(rows=1, cols=1)
    tabela.style = "Table Grid"
    celula = tabela.rows[0].cells[0]
    _sombrear(celula, CINZA_FUNDO)
    celula.text = ""
    p = celula.paragraphs[0]
    _run(p, "Como preencher: ", negrito=True, tamanho=8.5)
    _run(p, "os campos em ", tamanho=8.5)
    _run(p, "negrito", negrito=True, tamanho=8.5)
    _run(p, " já foram preenchidos pela contabilidade — apenas confira. Os "
            "campos marcados como ", tamanho=8.5)
    _run(p, MARCADOR_PENDENTE, negrito=True, tamanho=8.5, cor=COR_PENDENTE)
    _run(p, " dependem da sua decisão. Os demais são dados que você deve "
            "informar.", tamanho=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _rodape_assinatura(doc) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, "Declaração: ", negrito=True, tamanho=8)
    _run(p, "declaro que as informações prestadas neste formulário são "
            "verdadeiras e completas, e autorizo seu uso para os atos de "
            "constituição/alteração societária e cadastro nos órgãos "
            "competentes, nos termos da Lei nº 13.709/2018 (LGPD).", tamanho=8)

    doc.add_paragraph()
    tabela = doc.add_table(rows=2, cols=2)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, rotulo in enumerate(["Assinatura do titular / sócio administrador",
                                  "Local e data"]):
        c = tabela.cell(0, col)
        c.text = ""
        _run(c.paragraphs[0], "_" * 40, tamanho=10)
        c2 = tabela.cell(1, col)
        c2.text = ""
        _run(c2.paragraphs[0], rotulo, tamanho=8, cor=COR_CINZA)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"{CONTRATADA_NOME_FANTASIA} · {CONTRATADA_TELEFONE} · "
            f"{CONTRATADA_EMAIL}", tamanho=7.5, cor=COR_CINZA)


def gerar_formulario_abertura(
    perfil: str,
    dados_empresa: dict | None = None,
    dados_endereco: dict | None = None,
    socios: list[dict] | None = None,
    dados_desenquadramento: dict | None = None,
    minimo_socios: int = 2,
) -> bytes:
    """Monta o formulário e devolve os bytes do ``.docx``.

    ``perfil`` = ``"PF"`` (abertura nova), ``"MEI"`` (desenquadramento) ou
    ``"PJ"`` (alteração cadastral). Muda o título e quais seções entram.
    """
    perfil = perfil.upper()
    dados_empresa = dados_empresa or {}
    dados_endereco = dados_endereco or {}
    socios = socios or []

    titulos = {
        "MEI": ("FICHA CADASTRAL — DESENQUADRAMENTO DE MEI",
                "Migração de MEI para Microempresa (ME) no Simples Nacional"),
        "PF": ("FICHA CADASTRAL — ABERTURA DE EMPRESA",
               "Constituição de nova sociedade"),
        "PJ": ("FICHA CADASTRAL — ALTERAÇÃO CONTRATUAL",
               "Atualização de dados cadastrais e societários"),
    }
    titulo, subtitulo = titulos.get(perfil, titulos["PF"])

    doc = Document()
    _configurar_pagina(doc)
    _cabecalho(doc, titulo, f"{subtitulo} · emitida em {date.today():%d/%m/%Y}")
    _legenda(doc)

    if perfil == "MEI" and dados_desenquadramento:
        secao = _secao_desenquadramento(dados_desenquadramento)
        _titulo_secao(doc, secao.titulo)
        _tabela_campos(doc, secao.campos)

    for secao in (_secao_empresa(dados_empresa), _secao_endereco(dados_endereco)):
        _titulo_secao(doc, secao.titulo)
        _tabela_campos(doc, secao.campos)

    # Sempre imprime ao menos ``minimo_socios`` blocos: o cliente pode incluir
    # sócio que a contabilidade ainda não conhece.
    total_socios = max(len(socios), minimo_socios)
    for i in range(total_socios):
        dados = socios[i] if i < len(socios) else {}
        secao = _secao_socio(i + 1, dados)
        _titulo_secao(doc, secao.titulo)
        _tabela_campos(doc, secao.campos)

    _titulo_secao(doc, "Documentos a anexar")
    for item in _documentos(perfil):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _run(p, f"[   ]  {item}", tamanho=8.5)

    _rodape_assinatura(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _documentos(perfil: str) -> tuple[str, ...]:
    comuns = (
        "RG e CPF de todos os sócios (ou CNH)",
        "Comprovante de residência de cada sócio (últimos 3 meses)",
        "Comprovante de endereço do imóvel da empresa (IPTU ou conta de consumo)",
        "Contrato de locação do ponto, se alugado",
    )
    if perfil == "MEI":
        return (
            *comuns,
            "Certificado da Condição de MEI (CCMEI)",
            "Extrato do faturamento (relatório mensal de receitas do MEI)",
            "Últimos DAS-MEI pagos",
            "Declaração Anual do MEI (DASN-SIMEI) do último exercício",
            "Relatório de vendas por NCM/EAN do sistema de autoatendimento",
        )
    if perfil == "PJ":
        return (
            *comuns,
            "Contrato social e última alteração consolidada",
            "Cartão CNPJ atualizado",
            "Certificado digital e-CNPJ (A1 ou A3)",
        )
    return (
        *comuns,
        "Certidão de casamento, se casado(a)",
        "Consulta de viabilidade aprovada, se já solicitada",
    )


# --------------------------------------------------------------------------- #
# Ponte com os modelos do sistema                                              #
# --------------------------------------------------------------------------- #
def dados_de_contratante(contratante) -> tuple[dict, dict, list[dict]]:
    """Extrai do contratante o que já dá para preencher automaticamente.

    Devolve ``(dados_empresa, dados_endereco, socios)`` no formato que
    :func:`gerar_formulario_abertura` espera.
    """
    from ..core.pessoas import ContratantePJ

    endereco = contratante.endereco
    dados_endereco = {
        "logradouro": endereco.logradouro,
        "numero": endereco.numero,
        "complemento": endereco.complemento,
        "bairro": endereco.bairro,
        "municipio": endereco.municipio,
        "uf": endereco.uf,
        "cep": endereco.cep_formatado,
        "telefone": contratante.telefone,
        "email": contratante.email,
    }

    if isinstance(contratante, ContratantePJ):
        empresa = {
            "cnpj": formatar_cnpj(contratante.cnpj) if contratante.cnpj else "",
            "razao_social": contratante.razao_social,
            "nome_fantasia": contratante.nome_fantasia,
            "cnae": contratante.cnae_principal,
            "regime": contratante.regime,
        }
        rep = contratante.representante
        socios = [{
            "nome": rep.nome,
            "cpf": formatar_cpf(rep.cpf) if rep.cpf else "",
            "rg": rep.rg,
            "orgao": rep.orgao_emissor,
            "estado_civil": rep.estado_civil,
            "profissao": rep.profissao,
            "nacionalidade": rep.nacionalidade,
            "administrador": "Sim",
            **{k: dados_endereco[k] for k in
               ("logradouro", "numero", "bairro", "municipio", "uf", "cep")},
        }] if rep.nome else []
    else:
        empresa = {}
        socios = [{
            "nome": contratante.nome,
            "cpf": formatar_cpf(contratante.cpf) if contratante.cpf else "",
            "rg": contratante.rg,
            "orgao": contratante.orgao_emissor,
            "estado_civil": contratante.estado_civil,
            "profissao": contratante.profissao,
            "nacionalidade": contratante.nacionalidade,
            **{k: dados_endereco[k] for k in
               ("logradouro", "numero", "bairro", "municipio", "uf", "cep")},
            "telefone": contratante.telefone,
            "email": contratante.email,
        }] if contratante.nome else []

    return empresa, dados_endereco, socios


def capital_social_formatado(valor: float) -> tuple[str, str]:
    """``(numérico, por extenso)`` — usado quando o valor já foi definido."""
    if valor <= 0:
        return "", ""
    return moeda(valor), valor_extenso(valor)
